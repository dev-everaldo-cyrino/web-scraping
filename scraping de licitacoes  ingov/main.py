import click
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from time import sleep
import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import re




datas = datetime.datetime.now()
dic = {'1':'01', '2':'02', '3':'03', '4':'04', '5':'05', '6':'06', '7':'07', '8':'08', '9':'09', }
dia = str(datas.day)
mes = str(datas.month)
for x,y in dic.items():
    if dia == x:
        dia = y
    if mes == x:
        mes = y
data = dia+'/'+mes
data_completa = dia+'/'+mes+'/'+str(datas.year)

resdata = re.sub("/",".",data_completa)
document = Document()


lista_de_busca = []
link_para_scrapingg = []
unidade_item_lista = []
total_item_lista = []
palavras_item_lista = []
palavra_da_pesquisa = []

root = Tk()
root.geometry('700x600+0+0')
root.config(bg='SkyBlue4')
root.title('Pesquisa na 3ª seção do Diário Oficial da União')

data_atual = StringVar()

#radiobuttons
atual = Radiobutton(root,text='data atual',bg='SkyBlue4',font=('Arial',18),value='atual',variable=data_atual)
atual.place(x=10,y=60)
person = Radiobutton(root,text='personalizado',bg='SkyBlue4',font=('Arial',18),value='person',variable=data_atual)
person.place(x=10,y=100)

#spinbox da data personalizada
Label(root,text='Inicio',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=140)
inicio_dia = Spinbox(root, from_=1,to=31,width=3)
inicio_dia.place(x=10,y=170)
inicio_mes = Spinbox(root, from_=1,to=12,width=3)
inicio_mes.place(x=50,y=170)
inicio_ano = Spinbox(root, values=(2022),width=5)
inicio_ano.place(x=90,y=170)

Label(root,text='Fim',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=200)
fim_dia = Spinbox(root, from_=1,to=31,width=3)
fim_dia.place(x=10,y=230)
fim_mes = Spinbox(root, from_=1,to=12,width=3)
fim_mes.place(x=50,y=230)
fim_ano = Spinbox(root, values=(2022),width=5)
fim_ano.place(x=90,y=230)

#lista de itens
entrada_inicio = Entry(root,font='Times-New-Roman 16 bold',width="20")
entrada_inicio.place(x=200,y=60)
btn_add= Button(root,text='adicionar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:add(entrada_inicio.get())).place(x=200,y=90)

#treeview
tv = ttk.Treeview(root,columns=('chave','chave2'), show='headings')
tv.column('chave',minwidth=0,width=90)
tv.heading('chave', text='palavras chaves')
tv.column('chave2',minwidth=0,width=100)
tv.heading('chave2', text='')
tv.place(x=470,y=30)

#botoes abaixo da treeview
btn_buscar= Button(root,text='buscar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:buscar(data_atual.get())).place(x=470,y=260)


#conjunto busca personalizada
Label(root,text='busca personalizada por url',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=410)
urls = Entry(root,font='arial 16 bold',width=30)
urls.place(x=10,y=445)
btn_url= Button(root,text='busca de url',bg='DarkGoldenrod4',font='arial 14 bold',width='12',bd=7,command=lambda:scraping(urls.get())).place(x=10,y=485)

btn_fim= Button(root,text='gerar relatorio',bg='DarkGoldenrod4',font='arial 14 bold',width='12',bd=7,command=lambda:fim()).place(x=10,y=535)


#função de scraping por url
def fim():
    a = messagebox.askyesno(title='gerando documento ...',message='deseja finalizar ?')
    document.save("busca_por_url.docx")
    if a == True:
        root.destroy()   
def scraping(url_base):
    response = requests.get('{}'.format(url_base))
    sleep(0.5)
    content = response.content
    site = BeautifulSoup(content,'html.parser')
                
    hospedagem = site.findAll('p', attrs={'class': 'dou-paragraph'})
    cont ='\n\n'.join([detalhe.text for detalhe in hospedagem])
    sleep(0.7)
    dt =site.find('span', attrs={'class':'publicado-dou-data'})
    edicao = site.find('span', attrs={'class':'edicao-dou-data'})
    secao = site.find('span', attrs={'class':'secao-dou'})
    secao_detalhes = secao.text[0]+secao.text[1]+secao.text[2]+secao.text[3]+secao.text[4]+secao.text[5]+secao.text[6]+secao.text[7]
    pg =site.find('span', attrs={'class':'secao-dou-data'})
    orgao =site.find('span', attrs={'class':'orgao-dou-data'})
    tipo = site.findAll('p', attrs={'class': 'identifica'})
    tipo_detalhes ='\n\n'.join([tipo.text for tipo in tipo])
    dat =site.find('p', attrs={'class':'data'})
    assinatura = site.find('p', attrs={'class':'assina'})
    cargo =site.find('p', attrs={'class':'cargo'})

    sleep(0.7)   
    personalizada= url_base                       
    head = document.add_heading("Contratações pela Administração Pública",0)
    head.alignment = 0
    paragraph_1 = document.add_paragraph()
    paragraph_1.add_run('\nPublicado em: {} | Edição: {} | {} | página: {}'.format(dt.text,edicao.text,secao_detalhes,pg.text)).bold = True
    paragraph_2 = document.add_paragraph()
    paragraph_2.add_run("pesquisado por: ({})".format(personalizada)).bold = True
    paragraph_1 = document.add_paragraph('_________________________________________________________________________________________________________')
    paragraph_4 = document.add_paragraph()
    paragraph_4.add_run("{}".format(orgao.text)).bold = True
    paragraph_1 = document.add_paragraph("{}".format(tipo_detalhes))
    paragraph_1 = document.add_paragraph("{}".format(cont))
    if (dat):
        paragraph_1 = document.add_paragraph("{}".format(dat.text))
    if (assinatura):
        paragraph_1 = document.add_paragraph("{}".format(assinatura.text))
    if (cargo):
        paragraph_1 = document.add_paragraph("{}".format(cargo.text))
    messagebox.showinfo(title='info',message='sucesso!')
    urls.delete(0,END)
    
    

#função que adiciona valores na treeview
def add(valor_do_entre):
    if valor_do_entre == '' or valor_do_entre == ' ' or valor_do_entre == '  ' or valor_do_entre == '   ':
        messagebox.showerror(title='ERRO !',message='caixa de texto vasia !')
        entrada_inicio.delete(0,END)
    else:
        tv.insert("","end",values=(valor_do_entre))
        entrada_inicio.delete(0,END)
        lista_de_busca.append(valor_do_entre)
        

        
#função de buscar
def buscar(dtbusca):
    if lista_de_busca == []:
        messagebox.showerror(title='ERRO !',message='lista de busca vazia')
    else:
        
        res = []
        for lista in lista_de_busca:
            res.append(re.sub(" ", "+",lista))
        if dtbusca == 'atual':
            p = document.add_heading('Contratações no Diário Oficial da União do dia {}'.format(resdata), 0)
            p = document.add_paragraph()
            p.add_run('Data da publicação: {} | Data da pesquisa: {}'.format(resdata,resdata)).bold = True
            for x in res:
                url_base = 'https://www.in.gov.br/consulta/-/buscar/dou?q=%22{}%22&s=todos&exactDate=dia&sortType=0'.format(x)
                navegador = webdriver.Chrome()
                navegador.get('{}'.format(url_base))
                sleep(3)
                content = navegador.page_source
                site = BeautifulSoup(content,'html.parser')
                homes = site.findAll('div', attrs={'class': 'resultados-wrapper'})
                total_de_itens = 0
                unidade_de_itens = 0
                for home in homes:
                    selecao = home.find('li',attrs={'class': 'breadcrumb-item publication-info-marker'})
                    dts = selecao.text
                    dts = dts[16]+dts[17]+dts[18]+dts[19]+dts[20]
                    if dts == data:
                        total_de_itens += 1
                        homelink = home.find('h5')
                        s = homelink.text
                        s = "".join(s.split())
                        avs = s[0]+s[1]+s[2]+s[3]+s[4]
                        if len(s)> 11:
                            if s[7]+s[8] == 'ad' or s[7]+s[8] == 'AD' or s[7]+s[8] == 'Ad':
                                avs = s[7]+s[8] 
                            if len(s)> 22:
                                if s[19]+s[20]+s[21] == 'adj' or s[19]+s[20]+s[21] == 'ADJ' or s[19]+s[20]+s[21] == 'Adj':
                                    avs = s[19]+s[20]+s[21]          
                        else:
                            avs = s[0]+s[1]+s[2]+s[3]+s[4]
                        
                        if avs == 'aviso' or avs == 'AVISO' or avs == 'Aviso':
                            link = home.find('a')
                            sleep(1)
                            palavra_com_link = 'https://www.in.gov.br{}'.format(link['href'])
                            link_para_scrapingg.append(palavra_com_link) 
                            palavra_da_pesquisa.append(x)
                            unidade_de_itens +=1
                while True:
                    
                    if total_de_itens == 20 or total_de_itens == 40 or total_de_itens == 60 or total_de_itens == 80 or total_de_itens == 100 or total_de_itens == 120 or total_de_itens == 140:
                        try:
                            next_botao = navegador.find_element_by_id('rightArrow')
                            next_botao.click()
                            sleep(5)
                            content = navegador.page_source
                            site = BeautifulSoup(content,'html.parser')
                            homes = site.findAll('div', attrs={'class': 'resultados-wrapper'})
                            for home in homes:
                                selecao = home.find('li',attrs={'class': 'breadcrumb-item publication-info-marker'})
                                dts = selecao.text
                                dts = dts[16]+dts[17]+dts[18]+dts[19]+dts[20]
                                if dts == data:
                                    total_de_itens += 1
                                    homelink = home.find('h5')
                                    s = homelink.text
                                    s = "".join(s.split())
                                    avs = s[0]+s[1]+s[2]+s[3]+s[4]
                                    if len(s)> 11:
                                        if s[7]+s[8] == 'ad' or s[7]+s[8] == 'AD' or s[7]+s[8] == 'Ad':
                                            avs = s[7]+s[8] 
                                        if len(s)> 22:
                                            if s[19]+s[20]+s[21] == 'adj' or s[19]+s[20]+s[21] == 'ADJ' or s[19]+s[20]+s[21] == 'Adj':
                                                avs = s[19]+s[20]+s[21]          
                                    else:
                                        avs = s[0]+s[1]+s[2]+s[3]+s[4]
                                    
                                    if avs == 'aviso' or avs == 'AVISO' or avs == 'Aviso':
                                        link = home.find('a')
                                        sleep(1)
                                        palavra_com_link = 'https://www.in.gov.br{}'.format(link['href'])
                                        link_para_scrapingg.append(palavra_com_link) 
                                        palavra_da_pesquisa.append(x)
                                        unidade_de_itens +=1
                            
                            
                        except:
                            break   
                    else:
                        break
                unidade_item_lista.append(unidade_de_itens)
                total_item_lista.append(total_de_itens)
                palavras_item_lista.append(x.upper())
            for x in range(0,len(palavras_item_lista)):
                if total_item_lista[x]==0 or unidade_item_lista[x]==0:
                    p = document.add_paragraph()
                    p.add_run('{}  aviso de  {}  resultados   (0%) para o termo: ({}) '.format(unidade_item_lista[x],total_item_lista[x],palavras_item_lista[x])).bold = True
                else:
                    p = document.add_paragraph()
                    p.add_run('{}  aviso  de {}  resultados   ({:.2f}%) para o termo: ({}) '.format(unidade_item_lista[x],total_item_lista[x],((100*unidade_item_lista[x])/total_item_lista[x]),palavras_item_lista[x])).bold = True
            document.add_page_break()
            
            navegador.close()
            n = 0
            for x in link_para_scrapingg:
                response = requests.get('{}'.format(x))
                sleep(0.5)
                content = response.content
                site = BeautifulSoup(content,'html.parser')
                            
                hospedagem = site.findAll('p', attrs={'class': 'dou-paragraph'})
                cont ='\n\n'.join([detalhe.text for detalhe in hospedagem])
                sleep(0.7)
                dt =site.find('span', attrs={'class':'publicado-dou-data'})
                edicao = site.find('span', attrs={'class':'edicao-dou-data'})
                secao = site.find('span', attrs={'class':'secao-dou'})
                secao_detalhes = secao.text[0]+secao.text[1]+secao.text[2]+secao.text[3]+secao.text[4]+secao.text[5]+secao.text[6]+secao.text[7]
                pg =site.find('span', attrs={'class':'secao-dou-data'})
                orgao =site.find('span', attrs={'class':'orgao-dou-data'})
                tipo = site.findAll('p', attrs={'class': 'identifica'})
                tipo_detalhes ='\n\n'.join([tipo.text for tipo in tipo])
                dat =site.find('p', attrs={'class':'data'})
                assinatura = site.find('p', attrs={'class':'assina'})
                cargo =site.find('p', attrs={'class':'cargo'})

                sleep(0.7)   
                personalizada= url_base                       
                head = document.add_heading("Contratações publicadas no Diário Oficial da União do dia {}".format(data_completa),0)
                head.alignment = 0
                paragraph_1 = document.add_paragraph()
                paragraph_1.add_run('\nPublicado em: {} | Edição: {} | {} | página: {}'.format(dt.text,edicao.text,secao_detalhes,pg.text)).bold = True
                paragraph_2 = document.add_paragraph()
                paragraph_2.add_run("pesquisado por: ({})".format(palavra_da_pesquisa[n].upper())).bold = True
                n += 1
                paragraph_1 = document.add_paragraph('_________________________________________________________________________________________________________')
                paragraph_4 = document.add_paragraph()
                paragraph_4.add_run("{}".format(orgao.text)).bold = True
                paragraph_1 = document.add_paragraph("{}".format(tipo_detalhes))
                paragraph_1 = document.add_paragraph("{}".format(cont))
                if (dat):
                    paragraph_1 = document.add_paragraph("{}".format(dat.text))
                if (assinatura):
                    paragraph_1 = document.add_paragraph("{}".format(assinatura.text))
                if (cargo):
                    paragraph_1 = document.add_paragraph("{}".format(cargo.text))
                document.add_page_break()
           
            messagebox.showinfo(title='info',message='sucesso!')
            document.save('data_atual.docx')
            sleep(1)
            root.destroy()
                                    
            
            
        else:
            #inicio do spinbutton
            idia =str(inicio_dia.get())
            imes =str(inicio_mes.get())
            iano =str(inicio_ano.get())
            
            #fim do spinbutton
            fdia = str(fim_dia.get())
            fmes = str(fim_mes.get())
            fano = str(fim_ano.get())

            for x,y in dic.items():
                if idia == x:
                    idia = y
                if imes == x:
                    imes = y
                if fdia == x:
                    fdia = y
                if fmes == x:
                    fmes = y
            p = document.add_heading('Contratações no Diário Oficial da União do dia {}.{}.{}  a  {}.{}.{}'.format(idia,imes,iano,fdia,fmes,fano), 0)
            p = document.add_paragraph()
            p.add_run('Data da publicação: {}.{}.{}  a  {}.{}.{}| Data da pesquisa: {}'.format(idia,imes,iano,fdia,fmes,fano,resdata)).bold = True
            for x in res:
                #https://www.in.gov.br/consulta/-/buscar/dou?q=%22aterro+sanitario%22&s=todos&exactDate=personalizado&sortType=0&publishFrom=01-02-2022&publishTo=07-02-2022
                
                url_base = 'https://www.in.gov.br/consulta/-/buscar/dou?q=%22{}%22&s=todos&exactDate=personalizado&sortType=0&publishFrom={}-{}-{}&publishTo={}-{}-{}'.format(x,idia,imes,iano, fdia,fmes,fano)
                print(url_base)
                navegador = webdriver.Chrome()
                navegador.get('{}'.format(url_base))
                sleep(3)
                content = navegador.page_source
                site = BeautifulSoup(content,'html.parser')
                homes = site.findAll('div', attrs={'class': 'resultados-wrapper'})
                total_de_itens = 0
                unidade_de_itens = 0
                for home in homes:
                    selecao = home.find('li',attrs={'class': 'breadcrumb-item publication-info-marker'})
                    dts = selecao.text
                    dts = dts[16]+dts[17]+dts[18]+dts[19]+dts[20]
                
                    total_de_itens += 1
                    homelink = home.find('h5')
                    s = homelink.text
                    s = "".join(s.split())
                    avs = s[0]+s[1]+s[2]+s[3]+s[4]
                    if len(s)> 11:
                        if s[7]+s[8] == 'ad' or s[7]+s[8] == 'AD' or s[7]+s[8] == 'Ad':
                            avs = s[7]+s[8] 
                        if len(s)> 22:
                            if s[19]+s[20]+s[21] == 'adj' or s[19]+s[20]+s[21] == 'ADJ' or s[19]+s[20]+s[21] == 'Adj':
                                avs = s[19]+s[20]+s[21]          
                    else:
                        avs = s[0]+s[1]+s[2]+s[3]+s[4]
                    
                    if avs == 'aviso' or avs == 'AVISO' or avs == 'Aviso':
                        link = home.find('a')
                        sleep(1)
                        palavra_com_link = 'https://www.in.gov.br{}'.format(link['href'])
                        link_para_scrapingg.append(palavra_com_link) 
                        palavra_da_pesquisa.append(x)
                        unidade_de_itens +=1
                while True:
                    
                    if total_de_itens == 20 or total_de_itens == 40 or total_de_itens == 60 or total_de_itens == 80 or total_de_itens == 100 or total_de_itens == 120 or total_de_itens == 140 or total_de_itens == 140 or total_de_itens == 140 or total_de_itens == 140 or total_de_itens == 160 or total_de_itens == 180 or total_de_itens == 200 or total_de_itens == 220 or total_de_itens == 240 or total_de_itens == 260 or total_de_itens == 280 or total_de_itens == 300:
                        try:
                            next_botao = navegador.find_element_by_id('rightArrow')
                            next_botao.click()
                            sleep(5)
                            content = navegador.page_source
                            site = BeautifulSoup(content,'html.parser')
                            homes = site.findAll('div', attrs={'class': 'resultados-wrapper'})
                            for home in homes:
                                selecao = home.find('li',attrs={'class': 'breadcrumb-item publication-info-marker'})
                                dts = selecao.text
                                dts = dts[16]+dts[17]+dts[18]+dts[19]+dts[20]
                                
                                total_de_itens += 1
                                homelink = home.find('h5')
                                s = homelink.text
                                s = "".join(s.split())
                                avs = s[0]+s[1]+s[2]+s[3]+s[4]
                                if len(s)> 11:
                                    if s[7]+s[8] == 'ad' or s[7]+s[8] == 'AD' or s[7]+s[8] == 'Ad':
                                        avs = s[7]+s[8] 
                                    if len(s)> 22:
                                        if s[19]+s[20]+s[21] == 'adj' or s[19]+s[20]+s[21] == 'ADJ' or s[19]+s[20]+s[21] == 'Adj':
                                            avs = s[19]+s[20]+s[21]          
                                else:
                                    avs = s[0]+s[1]+s[2]+s[3]+s[4]
                                
                                if avs == 'aviso' or avs == 'AVISO' or avs == 'Aviso':
                                    link = home.find('a')
                                    sleep(1)
                                    palavra_com_link = 'https://www.in.gov.br{}'.format(link['href'])
                                    link_para_scrapingg.append(palavra_com_link) 
                                    palavra_da_pesquisa.append(x)
                                    unidade_de_itens +=1
                    
                    
                        except:
                            break
                    else:
                        break
                unidade_item_lista.append(unidade_de_itens)
                total_item_lista.append(total_de_itens)
                palavras_item_lista.append(x.upper())
            
            print(total_item_lista)
            print(unidade_item_lista)
            if total_item_lista[0] == 0 or unidade_item_lista[0]==0:
                p = document.add_paragraph()
                p.add_run('{}  aviso  de {}  resultados  (0%) para o termo: ({}) '.format(unidade_item_lista[-1],total_item_lista[-1],palavras_item_lista[-1])).bold = True
            else:
                p1 = document.add_paragraph()
                p1.add_run('{}  aviso  de {}  resultados  ({:.2f}%) para o termo: ({}) '.format(unidade_item_lista[-1],total_item_lista[-1],((100*unidade_item_lista[-1])/total_item_lista[-1]),palavras_item_lista[-1])).bold = True
            document.add_page_break()
            
            navegador.close()
            n = 0
            for x in link_para_scrapingg:
                response = requests.get('{}'.format(x))
                sleep(0.5)
                content = response.content
                site = BeautifulSoup(content,'html.parser')
                            
                hospedagem = site.findAll('p', attrs={'class': 'dou-paragraph'})
                cont ='\n\n'.join([detalhe.text for detalhe in hospedagem])
                sleep(0.7)
                dt =site.find('span', attrs={'class':'publicado-dou-data'})
                edicao = site.find('span', attrs={'class':'edicao-dou-data'})
                secao = site.find('span', attrs={'class':'secao-dou'})
                secao_detalhes = secao.text[0]+secao.text[1]+secao.text[2]+secao.text[3]+secao.text[4]+secao.text[5]+secao.text[6]+secao.text[7]
                pg =site.find('span', attrs={'class':'secao-dou-data'})
                orgao =site.find('span', attrs={'class':'orgao-dou-data'})
                tipo = site.findAll('p', attrs={'class': 'identifica'})
                tipo_detalhes ='\n\n'.join([tipo.text for tipo in tipo])
                dat =site.find('p', attrs={'class':'data'})
                assinatura = site.find('p', attrs={'class':'assina'})
                cargo =site.find('p', attrs={'class':'cargo'})

                sleep(0.7)   
                personalizada= url_base                       
                head = document.add_heading("Contratações publicadas no Diário Oficial da União do dia {}".format(dt.text),0)
                head.alignment = 0
                paragraph_1 = document.add_paragraph()
                paragraph_1.add_run('\nPublicado em: {} | Edição: {} | {} | página: {}'.format(dt.text,edicao.text,secao_detalhes,pg.text)).bold = True
                paragraph_2 = document.add_paragraph()
                paragraph_2.add_run("pesquisado por: ({})".format(palavra_da_pesquisa[n].upper())).bold = True
                n += 1
                paragraph_1 = document.add_paragraph('_________________________________________________________________________________________________________')
                paragraph_4 = document.add_paragraph()
                paragraph_4.add_run("{}".format(orgao.text)).bold = True
                paragraph_1 = document.add_paragraph("{}".format(tipo_detalhes))
                paragraph_1 = document.add_paragraph("{}".format(cont))
                if (dat):
                    paragraph_1 = document.add_paragraph("{}".format(dat.text))
                if (assinatura):
                    paragraph_1 = document.add_paragraph("{}".format(assinatura.text))
                if (cargo):
                    paragraph_1 = document.add_paragraph("{}".format(cargo.text))
                document.add_page_break()
           
            messagebox.showinfo(title='info',message='sucesso!')
            document.save('data_personalizada.docx')
            sleep(1)
            root.destroy()
        

root.mainloop()