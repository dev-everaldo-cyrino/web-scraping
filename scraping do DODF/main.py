from bs4 import BeautifulSoup
from selenium import webdriver
from time import sleep
from docx import Document
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
import pyautogui
import re
import requests




document = Document()
def busca_personalizada_por_data():
    pyautogui.PAUSE = 0.2

    lista_de_busca = []

    

    root1 = Toplevel()
    root1.geometry('700x600+0+0')
    root1.config(bg='SkyBlue4')
    root1.title('Pesquisa na 3ª seção do Diário Oficial da União')


    Label(root1,text='Para pesquisas anteriores \n a 01/05/2020 os filtros "tipo de seção", "tipo de jornal"\n e "onde pesquisar", não estão disponíveis',font=('arial', 8)).place(x=10,y=10)


    #spinbox da data personalizada
    Label(root1,text='DATA INICIAL',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=140)
    inicio_dia = Spinbox(root1, from_=1,to=31,width=3)
    inicio_dia.place(x=10,y=170)
    inicio_mes = Spinbox(root1, from_=1,to=12,width=3)
    inicio_mes.place(x=50,y=170)


    Label(root1,text='DATA FINAL',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=200)
    fim_dia = Spinbox(root1, from_=1,to=31,width=3)
    fim_dia.place(x=10,y=230)
    fim_mes = Spinbox(root1, from_=1,to=12,width=3)
    fim_mes.place(x=50,y=230)

    Label(root1,text='ano',bg='SkyBlue4',fg='black',font='arial 16 bold').place(x=10,y=260)
    fim_ano = Spinbox(root1,  from_=2020,to=2022 ,width=5)
    fim_ano.place(x=10,y=290)


    #lista de itens
    entrada_inicio = Entry(root1,font='Times-New-Roman 16 bold',width="20")
    entrada_inicio.place(x=200,y=135)
    btn_add= Button(root1,text='adicionar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:add(entrada_inicio.get())).place(x=200,y=170)

    #treeview
    tv = ttk.Treeview(root1,columns=('chave','chave2'), show='headings')
    tv.column('chave',minwidth=0,width=90)
    tv.heading('chave', text='palavras chaves')
    tv.column('chave2',minwidth=0,width=100)
    tv.heading('chave2', text='')
    tv.place(x=470,y=30)

    #botoes abaixo da treeview
    btn_buscar= Button(root1,text='buscar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:busca_url()).place(x=470,y=260)

    data_atual = StringVar()

    #radiobuttons
    aviso = Radiobutton(root1,text='AVISO',bg='SkyBlue4',font=('Arial',18),value='aviso',variable=data_atual)
    aviso.place(x=10,y=60)
    extrato = Radiobutton(root1,text='EXTRATO',bg='SkyBlue4',font=('Arial',18),value='extratos',variable=data_atual)
    extrato.place(x=10,y=100)


    #função que adiciona valores na treeview
    def add(valor_do_entre):
        if valor_do_entre == '' or valor_do_entre == ' ' or valor_do_entre == '  ' or valor_do_entre == '   ':
            messagebox.showerror(title='ERRO !',message='caixa de texto vasia !')
            entrada_inicio.delete(0,END)
        else:
            tv.insert("","end",values=(valor_do_entre))
            entrada_inicio.delete(0,END)
            lista_de_busca.append(valor_do_entre)
    
    
    def busca_url():
        p = document.add_heading('{}'.format(data_atual.get()), 0)
        url_base = 'https://www.dodf.df.gov.br/'
        navegador = webdriver.Chrome()
        navegador.get('{}'.format(url_base))
        navegador.find_element_by_id('filtroDePesquisa').click()
        sleep(1)
        navegador.find_element_by_xpath('//*[@id="formBuscaAvancada"]/div[2]/div/div/div[1]/div[2]/label[4]').click()    
        sleep(1)
        navegador.find_element_by_id('dtInicial').click()
        sleep(1)
        for x in range(1,6):
            pyautogui.press('backspace')
        dic = {'1':'01', '2':'02', '3':'03', '4':'04', '5':'05', '6':'06', '7':'07', '8':'08', '9':'09', }
        dia_inicio = str(inicio_dia.get())
        mes_inicio = str(inicio_mes.get())
        dia_fim = str(fim_dia.get())
        mes_fim = str(fim_mes.get())
        for x,y in dic.items():
            if dia_inicio == x:
                dia_inicio = y
            if mes_inicio == x:
                mes_inicio = y
            if dia_fim == x:
                dia_fim = y
            if mes_fim == x:
                mes_fim = y
    
        inicio = dia_inicio + '/' + mes_inicio
        fim = dia_fim + '/' + mes_fim
        pyautogui.write(inicio)
        pyautogui.press('tab')
        pyautogui.press('backspace')
        pyautogui.write(fim)
        pyautogui.press('tab')
        ano = int(fim_ano.get())
        ano = 2023-ano
        for x in range(0,ano):
            pyautogui.press('2')
        
        for listas in lista_de_busca:
            navegador.find_element_by_xpath('//*[@id="termo"]').click()
            for r in range(0,21):
                pyautogui.press('backspace')
            pyautogui.write(listas)
            navegador.find_element_by_xpath('//*[@id="btnProcurar"]').click()
            sleep(3)
            
            content = navegador.page_source
            site = BeautifulSoup(content,'html.parser')
            def scrap():
                atos = site.findAll('div', attrs = {'class': 'col-xl-12 mb-4 mt-4 itemMateria'})
                for ato in atos:
                    titulo = ato.find('a')
                    aviss = titulo.text[0] + titulo.text[1] +titulo.text[2] +titulo.text[3] 
                
                    if data_atual.get() == 'aviso':
                        if aviss == 'AVIS':
                            texto = ato.find('div',attrs={'class':'campoParteContexto'})
                            hora = ato.find('div',attrs={'class':'campoDtPublicacao'})
                            
                            paragraph_4 = document.add_paragraph()
                            paragraph_4.add_run("{}".format(titulo.text)).bold = True
                            paragraph_1 = document.add_paragraph("{}".format(titulo['href']))
                            paragraph_1 = document.add_paragraph("\n{}".format(texto.text))
                            paragraph_4 = document.add_paragraph()
                            paragraph_4.add_run("{}".format(hora.text)).bold = True
                            
                            
                           
                            document.add_page_break()
                    if data_atual.get() == 'extratos':
                        if aviss == 'EXTR':
                            texto = ato.find('div',attrs={'class':'campoParteContexto'})
                            hora = ato.find('div',attrs={'class':'campoDtPublicacao'})
                            
                            paragraph_4 = document.add_paragraph()
                            paragraph_4.add_run("{}".format(titulo.text)).bold = True
                            paragraph_1 = document.add_paragraph("{}".format(titulo['href']))
                            paragraph_1 = document.add_paragraph("\n{}".format(texto.text))
                            paragraph_4 = document.add_paragraph()
                            paragraph_4.add_run("{}".format(hora.text)).bold = True
                            
                            
                           
                            document.add_page_break()
                
                    
            
            scrap()
            pags = site.findAll('a',attrs={'target':'_self'})
            for pag in pags:
                if pag['href'] == 'javascript:void(0)':
                    pass
                else:
                    
                    response= requests.get('https://www.dodf.df.gov.br/{}'.format(pag['href']))            
                    sleep(0.5)
                    content = response.content
                    site = BeautifulSoup(content,'html.parser')
                    scrap()
        
        
        
        
    
        document.save('busca_personalizada.docx')
                    
        navegador.close()
        messagebox.showinfo(title='info',message='sucesso!')
        





root = Tk()
root.geometry('700x600+0+0')
root.config(bg='SkyBlue4')
root.title('Pesquisa no site DODF')

lista_de_busca=[]


#radiobuttons
data_atual = StringVar()

atual = Radiobutton(root,text='AVISO',bg='SkyBlue4',font=('Arial',16),value='aviso',variable=data_atual)
atual.place(x=10,y=60)
person = Radiobutton(root,text='EXTRATO',bg='SkyBlue4',font=('Arial',16),value='extrato',variable=data_atual)
person.place(x=10,y=100)
person = Radiobutton(root,text='AVISO E EXTRATO',bg='SkyBlue4',font=('Arial',16),value='aviso_extrato',variable=data_atual)
person.place(x=10,y=140)


#lista de itens
entrada_inicio = Entry(root,font='Times-New-Roman 16 bold',width="20")
entrada_inicio.place(x=250,y=265)
btn_add= Button(root,text='adicionar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:add(entrada_inicio.get())).place(x=300,y=300)

#treeview
tv = ttk.Treeview(root,columns=('chave','chave2'), show='headings')
tv.column('chave',minwidth=0,width=90)
tv.heading('chave', text='palavras chaves')
tv.column('chave2',minwidth=0,width=200)
tv.heading('chave2', text='')
tv.place(x=230,y=30)

#botoes busca
btn_buscar= Button(root,text='buscar',bg='DarkGoldenrod4',font='arial 14 bold',width='10',bd=7,command=lambda:buscar(data_atual.get())).place(x=300,y=360)

#botão personalizado
btn_busca_personalizado= Button(root,text='busca personalizada',bg='DarkGoldenrod4',font='arial 14 bold',width='20',bd=7,command=lambda:busca_personalizada_por_data()).place(x=10,y=360)

#função que adiciona valores na treeview
def add(valor_do_entre):
    if valor_do_entre == '' or valor_do_entre == ' ' or valor_do_entre == '  ' or valor_do_entre == '   ':
        messagebox.showerror(title='ERRO !',message='caixa de texto vasia !')
        entrada_inicio.delete(0,END)
    else:
        tv.insert("","end",values=(valor_do_entre))
        entrada_inicio.delete(0,END)
        lista_de_busca.append(valor_do_entre)


#função de buscar
Document = Document()
def buscar(tipo_do_ato):
    
    if lista_de_busca == []:
        messagebox.showerror(title='ERRO !',message='lista de busca vazia')
    else:
        #entrando no site e realizando o carregamento do react
        messagebox.showwarning(title='ATENÇÃO',message='atenção não clique em nada durante a execução do software. obrigado!.')
        url_base = 'https://www.dodf.df.gov.br/'
        navegador = webdriver.Chrome()
        navegador.get('{}'.format(url_base))
        sleep(3)
                    
        btn_secao3 = navegador.find_element_by_xpath('/html/body/div[3]/div[2]/div[2]/div/div[3]/div/div[2]/a').click()
        sleep(3)
        busca_aviso = BeautifulSoup(navegador.page_source, 'html.parser')
        principal = busca_aviso.find('select', attrs={'class':'slc-tp-ato'}).findAll('option')
        num=1
        for w in principal:
            sleep(0.5)
            if num ==2 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[2]'
            elif num ==3 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[3]'
            elif num ==4 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[4]'
            elif num ==5 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[5]'
            elif num ==6 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[6]'
            elif num ==7 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[7]'
            elif num ==8 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[8]'
            elif num ==9 and w.text == 'Aviso':
                pathx= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[9]'  
            
            if num ==2 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[2]'
            elif num ==3 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[3]'
            elif num ==4 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[4]'
            elif num ==5 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[5]'
            elif num ==6 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[6]'
            elif num ==7 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[7]'
            elif num ==8 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[8]'
            elif num ==9 and w.text == 'Extrato':
                pathy= '/html/body/div[4]/div[1]/div/div[2]/div[2]/select/option[9]'  
                
            num +=1
        
        
        
        #funcao anonima que realiza o scraping
        def busca():
            site = BeautifulSoup(navegador.page_source, 'html.parser')
            
            try:
                orgao = site.find('div', attrs={'class': 'demandante-orgao org-1'})          
                orgao = orgao.text
                conteudos = site.findAll('div' , attrs={'class': 'col-xl-12 mb-4 mt-4 itemMateria'})
                for conteudo in conteudos:
                    titulo = conteudo.find('p' , attrs={'class': 'titulo'})
                    aviso = titulo.text
                    aviso_final = aviso[0]+aviso[1]+aviso[2]+aviso[3]+aviso[4]+aviso[5]+aviso[6]+aviso[7]+aviso[8]+aviso[9]+aviso[10]
                    
                    if aviso_final == 'AVISO DE HO':
                        print('')
                    else:  
                        texto = conteudo.find('p' , attrs={'style': 'text-align:justify;'})
                        valor = 0
                        if (texto != None):
                            texto = texto.text
                            texto = texto.upper()
                            for x in lista_de_busca:
                                res=re.search(x.upper(),texto)
                                if (res !=None):
                                    valor =1
                                else:
                                    pass
                                
                            sleep(0.5)
                            if valor ==1:                          
                                
                                licitacao = conteudo.findAll('p',attrs={'style':'text-align:center;'})
                                texto = conteudo.findAll('p' , attrs={'style': 'text-align:justify;'})
                                head = Document.add_heading("{} - DF".format(orgao.upper()),1)
                                head.alignment = 1
                                paragraph1 = Document.add_paragraph()          
                                paragraph1.add_run("{}".format(titulo.text)).bold=True
                                paragraph1.alignment = 1
                                
                                if (licitacao):
                                    paragraph1 = Document.add_paragraph()          
                                    paragraph1.add_run("{}".format(licitacao[0].text)).bold=True
                                    paragraph1.alignment = 1
                                sleep(1)
                                
                                for y in range(0,len(texto)):
                                    paragraph1 = Document.add_paragraph(texto[y].text)
                                
                                if (licitacao):
                                    for x in range(1,len(licitacao)):
                                        paragraph1 = Document.add_paragraph("{}".format(licitacao[x].text))          
                                        paragraph1.alignment = 1
                                sleep(1)
                                Document.add_page_break()
                                
                                    
                    
            except:
                pass
        
                
        
            
            


        #comando de navegação entre os orgaos, cada orgao carrega um react e chama a funcao anonima para realizar o scraping da pagina
        
        
        
        #busca de avisos e extratos
        if tipo_do_ato == 'aviso_extrato':
            sleep(1)
            btn_aviso = navegador.find_element_by_xpath(pathx).click()
            sleep(2)
            btn_ato = navegador.find_element_by_class_name('slc-orgao-demandante').click()
            secundario = busca_aviso.find('select', attrs={'class':'slc-orgao-demandante'}).findAll('option')
            head = Document.add_heading("AVISOS",0)
            head.alignment = 1
            pyautogui.press('enter')
            for ww in range(1,len(secundario)):
                sleep(1)
                wwvalor = secundario[ww].text
                pyautogui.press(wwvalor[0])
                busca()
            sleep(2)
            
            
            
            sleep(1)
            btn_aviso = navegador.find_element_by_xpath(pathy).click()
            sleep(2)
            btn_ato = navegador.find_element_by_class_name('slc-orgao-demandante').click()
            secundario = busca_aviso.find('select', attrs={'class':'slc-orgao-demandante'}).findAll('option')
            head = Document.add_heading("EXTRATOS",0)
            head.alignment = 1
            pyautogui.press('e')
            pyautogui.press('enter')
            for ww in range(1,len(secundario)):
                sleep(1)
                wwvalor = secundario[ww].text
                pyautogui.press(wwvalor[0])
                busca()
            sleep(2)            
            Document.save("DODF_aviso_extrato.docx")
            navegador.close()
            messagebox.showinfo(title='info',message='sucesso! , pode voltar a usar o computador')
            root.destroy()
            
        
        #busca de avisos    
        elif tipo_do_ato == 'aviso':
            sleep(1)
            btn_aviso = navegador.find_element_by_xpath(pathx).click()
            sleep(2)
            btn_ato = navegador.find_element_by_class_name('slc-orgao-demandante').click()
            secundario = busca_aviso.find('select', attrs={'class':'slc-orgao-demandante'}).findAll('option')
            head = Document.add_heading("AVISOS",0)
            head.alignment = 1
            pyautogui.press('enter')
            for ww in range(1,len(secundario)):
                sleep(1)
                wwvalor = secundario[ww].text
                pyautogui.press(wwvalor[0])
                busca()
            sleep(2)
            Document.save("DODF_aviso.docx")
            navegador.close()
            messagebox.showinfo(title='info',message='sucesso! , pode voltar a usar o computador')
            root.destroy()
        
        
        #busca de extratos
        elif tipo_do_ato == 'extrato':
            sleep(1)
            btn_aviso = navegador.find_element_by_xpath(pathy).click()
            sleep(2)
            btn_ato = navegador.find_element_by_class_name('slc-orgao-demandante').click()
            secundario = busca_aviso.find('select', attrs={'class':'slc-orgao-demandante'}).findAll('option')
            head = Document.add_heading("EXTRATOS",0)
            head.alignment = 1
            pyautogui.press('enter')
            for ww in range(1,len(secundario)):
                sleep(1)
                wwvalor = secundario[ww].text
                pyautogui.press(wwvalor[0])
                busca()
            sleep(2)
            Document.save("DODF_extrato.docx")
            navegador.close()
            messagebox.showinfo(title='info',message='sucesso! , pode voltar a usar o computador')
            root.destroy()
            



root.mainloop()